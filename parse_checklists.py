import os
import json
import re
from docx import Document

def parse_standard_checklist(docx_path, json_path, system_type):
    """Parse standard ISO checklist docx into structured JSON."""
    doc = Document(docx_path)
    checklist = []
    current_clause = None
    current_category = None

    # Question prefixes to identify requirement rows
    question_prefixes = ["Have you", "Do you", "Can your", "Has ", "Is there", "Are there", "Does the", "Did you"]

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            if len(row.cells) == 0:
                continue
                
            first_cell_text = row.cells[0].text.strip()
            
            if not first_cell_text:
                continue
                
            # Check if this is a clause header
            if first_cell_text.startswith("Clause"):
                parts = first_cell_text.split(" ", 2)
                if len(parts) >= 3:
                    current_clause = parts[1]
                    current_category = parts[2]
                else:
                    current_clause = parts[1] if len(parts) >= 2 else ""
                    current_category = ""
                continue
                
            # Check if this is a requirement/question
            if any(first_cell_text.startswith(prefix) for prefix in question_prefixes):
                requirement_text = first_cell_text
                checklist.append({
                    "clause": current_clause,
                    "requirement": requirement_text,
                    "category": current_category,
                    "system": system_type
                })

    os.makedirs(os.path.dirname(json_path), exist_ok=True)

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(checklist, f, indent=2, ensure_ascii=False)

    print(f"✅ Saved {json_path} with {len(checklist)} entries")

def parse_stage1_checklist(docx_path, json_path, system_type):
    """Parse Stage 1 checklist with different table structure."""
    doc = Document(docx_path)
    checklist = []
    current_element = None
    current_clause = None

    # Question prefixes for Stage 1
    question_prefixes = ["Have you", "Do you", "Has ", "Does the", "Is this", "Are there"]
    
    # Element patterns (like "4.3 Scope of the Management System")
    element_pattern = r"^(\d+\.\d+)\s+(.+)"

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 1:
                continue
                
            element_cell = row.cells[0].text.strip()
            
            if not element_cell:
                continue

            # Check if this is an element header (like "4.3 Scope of the Management System")
            element_match = re.match(element_pattern, element_cell)
            if element_match:
                current_clause = element_match.group(1)  # e.g., "4.3"
                current_element = element_match.group(2)  # e.g., "Scope of the Management System"
                continue
            
            # Check if this is a main requirement question
            if any(element_cell.startswith(prefix) for prefix in question_prefixes):
                # For main questions, use the current element as category
                checklist.append({
                    "clause": current_clause,
                    "requirement": element_cell,
                    "category": current_element,
                    "system": system_type
                })
            
            # Check for sub-questions (starting with dash)
            elif element_cell.startswith("- ") and current_clause:
                # For sub-questions, prepend the main element for context
                full_requirement = f"{current_element}: {element_cell[2:]}"
                checklist.append({
                    "clause": current_clause,
                    "requirement": full_requirement,
                    "category": current_element,
                    "system": system_type
                })
            
            # Check for policy sub-questions (like "5.2 Have you established...")
            elif re.match(r"^>\s*\d+\.\d+\s+", element_cell):
                # This handles lines like "> 5.2 Have you established an Environmental Policy?"
                policy_match = re.match(r">\s*(\d+\.\d+)\s+(.+)", element_cell)
                if policy_match:
                    current_clause = policy_match.group(1)
                    current_element = f"Policy - {policy_match.group(2)}"
                    checklist.append({
                        "clause": current_clause,
                        "requirement": policy_match.group(2),
                        "category": "Policy",
                        "system": system_type
                    })

    # Also check paragraphs for any missed requirements
    for para in doc.paragraphs:
        text = para.text.strip()
        if any(text.startswith(prefix) for prefix in question_prefixes):
            checklist.append({
                "clause": "General",
                "requirement": text,
                "category": "General",
                "system": system_type
            })

    os.makedirs(os.path.dirname(json_path), exist_ok=True)

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(checklist, f, indent=2, ensure_ascii=False)

    print(f"✅ Saved {json_path} with {len(checklist)} Stage 1 entries")

def parse_checklist(docx_path, json_path, system_type):
    """Main parsing function that routes to appropriate parser."""
    if system_type == "Stage1":
        parse_stage1_checklist(docx_path, json_path, system_type)
    else:
        parse_standard_checklist(docx_path, json_path, system_type)

if __name__ == "__main__":
    base = os.path.dirname(__file__)
    docs = os.path.join(base, "docs")
    out = os.path.join(base, "checklists")

    files = [
        ("ATLAS Certification Checklist 9001-2015 (25).docx", "checklist_9001.json", "QMS"),
        ("ATLAS Certification Checklist 14001-2015 (25).docx", "checklist_14001.json", "EMS"), 
        ("ATLAS Certification Checklist 45001-2018 (25).docx", "checklist_45001.json", "OHSMS"),
        ("stage 1 checklist.docx", "checklist_stage1.json", "Stage1")
    ]

    for docx_name, json_name, sys_type in files:
        docx_path = os.path.join(docs, docx_name)
        json_path = os.path.join(out, json_name)

        if os.path.exists(docx_path):
            print(f"📋 Processing: {docx_name} -> {sys_type}")
            parse_checklist(docx_path, json_path, sys_type)
        else:
            print(f"⚠️ Missing: {docx_path}")
