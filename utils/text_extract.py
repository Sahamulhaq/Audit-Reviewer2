# utils/text_extract.py
import io
import re
import pdfplumber
from docx import Document

def clean_text(text: str) -> str:
    """Clean extracted text by removing excess whitespace and normalizing."""
    if not text:
        return ""
    # Replace multiple whitespace characters with single space
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF with fallback to simple parsing if pdfplumber fails."""
    text = ""
    
    # Try pdfplumber first (better for text-based PDFs)
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                # Try to extract text
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages_text.append(page_text)
                else:
                    # If no text, try to extract from tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            for row in table:
                                if row:
                                    row_text = ' '.join([str(cell) for cell in row if cell])
                                    if row_text.strip():
                                        pages_text.append(row_text)
            text = "\n".join(pages_text)
    except Exception as e:
        print(f"PDF extraction with pdfplumber failed: {e}")
        # Fallback: try simple text extraction with PyPDF2
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages_text.append(page_text)
            text = "\n".join(pages_text)
        except Exception as e2:
            print(f"PDF extraction fallback also failed: {e2}")
            text = ""
    
    cleaned = clean_text(text)
    
    # If we got very little text, the PDF might be image-based
    if len(cleaned) < 100 and len(file_bytes) > 10000:  # Small text but large file = likely image PDF
        print("PDF appears to be image-based, consider adding OCR functionality")
    
    return cleaned

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX files including tables."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract from paragraphs
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        tables_text.append(cell.text)
        
        all_text = paras + tables_text
        return clean_text("\n".join(all_text))
        
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""

# Optional: OCR functionality for image-based PDFs (uncomment if needed)
"""
def extract_text_from_pdf_with_ocr(file_bytes: bytes) -> str:
    \"\"\"Extract text from image-based PDFs using OCR.\"\"\"
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        
        images = convert_from_bytes(file_bytes, dpi=200)
        text = "\n".join([pytesseract.image_to_string(img) for img in images])
        return clean_text(text)
    except ImportError:
        print("OCR dependencies not installed. Install with: pip install pdf2image pytesseract")
        return ""
    except Exception as e:
        print(f"OCR extraction failed: {e}")
        return ""
"""
